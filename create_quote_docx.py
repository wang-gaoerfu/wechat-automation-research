from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.font.size = Pt(11)

BLUE = RGBColor(26, 95, 122)
RED = RGBColor(192, 57, 43)
GREEN = RGBColor(39, 174, 96)

# 封面
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title_run = title.add_run('微信自动化运营方案')
title_run.font.size = Pt(32)
title_run.font.bold = True
title_run.font.color.rgb = BLUE
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run('报 价 单')
subtitle_run.font.size = Pt(24)
subtitle_run.font.bold = True
subtitle_run.font.color.rgb = RED
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('文档版本：v2.0\n').font.size = Pt(11)
info.add_run('编制日期：2026年5月25日\n').font.size = Pt(11)
info.add_run('方案设计者：墨飞科技').font.size = Pt(11)

doc.add_page_break()

# 一、方案说明
h = doc.add_heading('一、方案说明', level=1)
h.runs[0].font.color.rgb = BLUE

doc.add_paragraph('本报价适用于小微型企业的微信自动化运营系统建设。采用AI编程方式开发，标准化交付流程，首家客户低利润甚至保本，后续客户成本逐步降低。')

p = doc.add_paragraph()
run = p.add_run('定价原则：')
run.bold = True

doc.add_paragraph('首单作为案例积累，低利润', style='List Bullet')
doc.add_paragraph('标准化复刻，后续客户逐步降价', style='List Bullet')
doc.add_paragraph('批量实施，价格更优', style='List Bullet')

doc.add_page_break()

# 二、产品版本
h = doc.add_heading('二、产品版本', level=1)
h.runs[0].font.color.rgb = BLUE

# 企业微信基础版
doc.add_heading('2.1 企业微信 - 基础版', level=2)
doc.add_paragraph('适合只需要核心功能（自动回复+定时群发）的小微企业。')

table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '功能模块'
hdr[1].text = '说明'
hdr[2].text = '工作量'
for cell in hdr:
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

data = [
    ('消息自动回复', '关键词触发，自动回复', '3天'),
    ('定时群发', '定时向客户推送内容', '2天'),
    ('客户标签管理', '自动打标签', '1天'),
    ('系统部署调试', '服务器部署', '1天'),
    ('培训支持', '操作培训', '1天'),
    ('合计', '', '8天'),
]
for i, (m, d, t) in enumerate(data, 1):
    table.rows[i].cells[0].text = m
    table.rows[i].cells[1].text = d
    table.rows[i].cells[2].text = t
    if i == 6:
        for cell in table.rows[i].cells:
            cell.paragraphs[0].runs[0].font.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('标准报价：¥8,000')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = BLUE

doc.add_paragraph()

# 企业微信完整版
doc.add_heading('2.2 企业微信 - 完整版', level=2)
doc.add_paragraph('适合需要全部功能的成长型企业。')

table2 = doc.add_table(rows=9, cols=3)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
hdr2[0].text = '功能模块'
hdr2[1].text = '说明'
hdr2[2].text = '工作量'
for cell in hdr2:
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

data2 = [
    ('基础版全部功能', '', '8天'),
    ('自动加好友', '生成获客二维码', '2天'),
    ('自动欢迎语', '新客添加自动回复', '1天'),
    ('智能助手', 'Agent+MCP，自然语言控制', '3天'),
    ('知识库RAG', 'ChromaDB语义搜索', '2天'),
    ('基础设施', '日志、监控、备份', '1天'),
    ('合计', '', '17天'),
]
for i, (m, d, t) in enumerate(data2, 1):
    table2.rows[i].cells[0].text = m
    table2.rows[i].cells[1].text = d
    table2.rows[i].cells[2].text = t
    if i == 7:
        for cell in table2.rows[i].cells:
            cell.paragraphs[0].runs[0].font.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('标准报价：¥13,600')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = BLUE

doc.add_paragraph()

# 个人微信方案
doc.add_heading('2.3 个人微信方案', level=2)
doc.add_paragraph('备选方案，适合无企业微信的客户。')

table3 = doc.add_table(rows=10, cols=3)
table3.style = 'Table Grid'
hdr3 = table3.rows[0].cells
hdr3[0].text = '功能模块'
hdr3[1].text = '说明'
hdr3[2].text = '工作量'
for cell in hdr3:
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

data3 = [
    ('基础框架', '项目结构、数据库', '1.5天'),
    ('WeChatFerry对接', 'RPC通信、消息订阅', '2.5天'),
    ('防风控系统', '频率限制、内容多样化', '4天'),
    ('客户管理', '标签、好友、存档', '2天'),
    ('消息功能', '关键词回复、定时群发', '3天'),
    ('智能助手', 'Agent+MCP架构', '3天'),
    ('版本适配', '微信版本检测', '2天'),
    ('基础设施', '日志、监控、重连', '3天'),
    ('合计', '', '21天'),
]
for i, (m, d, t) in enumerate(data3, 1):
    table3.rows[i].cells[0].text = m
    table3.rows[i].cells[1].text = d
    table3.rows[i].cells[2].text = t
    if i == 9:
        for cell in table3.rows[i].cells:
            cell.paragraphs[0].runs[0].font.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('标准报价：¥16,800')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = BLUE

doc.add_page_break()

# 三、客户阶梯报价
h = doc.add_heading('三、客户阶梯报价', level=1)
h.runs[0].font.color.rgb = BLUE

table4 = doc.add_table(rows=4, cols=4)
table4.style = 'Table Grid'
hdr4 = table4.rows[0].cells
hdr4[0].text = '客户顺序'
hdr4[1].text = '企业微信基础版'
hdr4[2].text = '企业微信完整版'
hdr4[3].text = '个人微信'
for cell in hdr4:
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

data4 = [
    ('首单（案例）', '¥4,000', '¥6,800', '¥8,400'),
    ('第2-5家', '¥6,000', '¥10,200', '¥12,600'),
    ('第6家起', '¥5,000', '¥8,500', '¥10,500'),
]
for i, row_data in enumerate(data4, 1):
    for j, val in enumerate(row_data):
        table4.rows[i].cells[j].text = val

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('说明：')
run.bold = True

doc.add_paragraph('首单低利润作为案例，用于建立口碑和案例', style='List Bullet')
doc.add_paragraph('后续随标准化程度提高，复刻成本降低，价格更优惠', style='List Bullet')
doc.add_paragraph('批量客户可单独议价', style='List Bullet')

doc.add_page_break()

# 四、增值服务
h = doc.add_heading('四、增值服务', level=1)
h.runs[0].font.color.rgb = BLUE

table5 = doc.add_table(rows=7, cols=3)
table5.style = 'Table Grid'
hdr5 = table5.rows[0].cells
hdr5[0].text = '服务项目'
hdr5[1].text = '说明'
hdr5[2].text = '报价'
for cell in hdr5:
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

data5 = [
    ('公众号对接', '公众号自动回复、内容推送', '¥3,000'),
    ('CRM对接', '与企业现有CRM系统对接', '¥5,000'),
    ('多企业支持', '同时管理多个企业微信', '¥2,000/企业'),
    ('年度维护', 'bug修复、版本更新、答疑', '¥3,000/年'),
    ('私有化部署', '部署到客户指定服务器', '¥1,500/次'),
    ('功能定制', '超出标准功能的定制开发', '¥800/人天'),
]
for i, (item, desc, price) in enumerate(data5, 1):
    table5.rows[i].cells[0].text = item
    table5.rows[i].cells[1].text = desc
    table5.rows[i].cells[2].text = price

doc.add_page_break()

# 五、交付物
h = doc.add_heading('五、交付物', level=1)
h.runs[0].font.color.rgb = BLUE

table6 = doc.add_table(rows=5, cols=2)
table6.style = 'Table Grid'
hdr6 = table6.rows[0].cells
hdr6[0].text = '类型'
hdr6[1].text = '内容'
for cell in hdr6:
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

data6 = [
    ('源码', '微信自动化系统完整源码'),
    ('文档', '部署手册、用户指南、技术文档'),
    ('培训', '线上/线下操作培训'),
    ('维护', '1年免费bug修复'),
]
for i, (t, c) in enumerate(data6, 1):
    table6.rows[i].cells[0].text = t
    table6.rows[i].cells[1].text = c

doc.add_page_break()

# 六、FAQ
h = doc.add_heading('六、常见问题', level=1)
h.runs[0].font.color.rgb = BLUE

faqs = [
    ('Q：为什么首单这么便宜？',
     'A：首单作为案例积累，我们将以保本价格甚至微利完成项目，用于建立口碑和案例。后续客户将按标准报价执行。'),
    ('Q：复刻项目为什么更便宜？',
     'A：得益于标准化设计，首家交付后，后续项目可复用大部分代码和配置，复刻成本大幅降低。'),
    ('Q：付款方式是怎样的？',
     'A：合同签订付30%首付款，系统验收后付70%尾款。'),
    ('Q：开发过程中需要客户配合什么？',
     'A：提供企业微信账号权限、明确需求范围、配合验收测试即可。'),
]

for q, a in faqs:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    run.font.color.rgb = BLUE
    doc.add_paragraph(a)

doc.add_page_break()

# 七、联系信息
h = doc.add_heading('七、联系信息', level=1)
h.runs[0].font.color.rgb = BLUE

table7 = doc.add_table(rows=5, cols=2)
table7.style = 'Table Grid'
data7 = [
    ('项目', '内容'),
    ('联系人', '[您的姓名]'),
    ('电话', '[您的电话]'),
    ('邮箱', '[您的邮箱]'),
    ('微信', '[您的微信]'),
]
for i, (k, v) in enumerate(data7):
    table7.rows[i].cells[0].text = k
    table7.rows[i].cells[1].text = v
    if i == 0:
        table7.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        table7.rows[i].cells[1].paragraphs[0].runs[0].font.bold = True
        table7.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        table7.rows[i].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('方案设计者：墨飞科技 | 本报价单最终解释权归墨飞科技所有')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(153, 153, 153)

doc.save('D:/my_project/ai_projects/wechat-automation-research/微信自动化方案-报价.docx')
print('报价Word文档已生成！')
