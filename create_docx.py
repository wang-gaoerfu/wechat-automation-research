from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.font.size = Pt(11)

# 封面
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title_run = title.add_run('微信自动化运营方案')
title_run.font.size = Pt(32)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(26, 95, 122)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run('企业微信 · 个人微信 · 智能助手')
subtitle_run.font.size = Pt(16)
subtitle_run.font.color.rgb = RGBColor(102, 102, 102)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('文档版本：v1.1\n').font.size = Pt(11)
info.add_run('编制日期：2026年5月25日\n').font.size = Pt(11)
info.add_run('目标读者：企业客户').font.size = Pt(11)

doc.add_page_break()

# 一、方案概述
h = doc.add_heading('一、方案概述', level=1)
h.runs[0].font.color.rgb = RGBColor(26, 95, 122)

p = doc.add_paragraph()
run = p.add_run('核心价值：')
run.bold = True
run.font.color.rgb = RGBColor(26, 95, 122)

doc.add_paragraph('节省人工操作时间，提升响应速度', style='List Bullet')
doc.add_paragraph('规范化客户管理，避免遗漏', style='List Bullet')
doc.add_paragraph('智能化内容推送，提高客户活跃度', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('微信自动化运营平台是一套帮助企业高效管理微信客户的解决方案。通过技术手段实现客户添加、消息回复、内容推送的自动化，大幅提升运营效率。')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('本方案提供两种实现路径：').bold = True
doc.add_paragraph('企业微信方案（推荐）：基于腾讯官方API，稳定合规', style='List Bullet')
doc.add_paragraph('个人微信方案（备选）：基于第三方软件，功能灵活', style='List Bullet')

doc.add_page_break()

# 二、企业微信方案
h = doc.add_heading('二、企业微信方案（推荐）', level=1)
h.runs[0].font.color.rgb = RGBColor(26, 95, 122)

doc.add_heading('2.1 什么是企业微信？', level=2)
doc.add_paragraph('企业微信是腾讯官方推出的企业级通讯工具，与个人微信互通。它提供官方API接口，支持自动化管理，且不会封号。')

doc.add_heading('2.2 核心功能', level=2)

table = doc.add_table(rows=8, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '功能模块'
hdr[1].text = '说明'
hdr[2].text = '示例'
for cell in hdr:
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

data = [
    ('自动加好友', '生成获客二维码，客户扫码自动添加', '线上推广、线下活动快速积累客户'),
    ('自动欢迎语', '新客户添加后立即发送欢迎消息', '"您好！很高兴认识您，有任何问题随时联系我"'),
    ('自动标签', '根据来源自动打标签分类', '来自不同活动的客户自动归类'),
    ('关键词回复', '客户发送特定关键词，自动回复对应内容', '客户发"价格"→自动回复价目表'),
    ('定时群发', '定时向客户推送内容', '每周一早上9点推送本周精选'),
    ('群内自动回复', '客户在群里@小助手，自动回复', 'FAQ、常见问题24小时自动解答'),
    ('智能助手', '用自然语言控制自动化流程', '"给所有VIP客户发送新品通知"'),
]
for i, (f, d, e) in enumerate(data, 1):
    table.rows[i].cells[0].text = f
    table.rows[i].cells[1].text = d
    table.rows[i].cells[2].text = e

doc.add_paragraph()
doc.add_heading('2.3 优势特点', level=2)
doc.add_paragraph('官方合规：使用腾讯官方API，永不封号', style='List Bullet')
doc.add_paragraph('稳定可靠：企业级服务，7×24小时运行', style='List Bullet')
doc.add_paragraph('操作简单：配置好后自动运行，大部分操作无需人工介入', style='List Bullet')
doc.add_paragraph('数据安全：客户资料本地存储，安全可控', style='List Bullet')

doc.add_heading('2.4 使用前提', level=2)
doc.add_paragraph('已开通企业微信', style='List Bullet')
doc.add_paragraph('如需"自动加好友"功能，需要完成企业认证（个人工商户也可）', style='List Bullet')

doc.add_page_break()

# 三、个人微信方案
h = doc.add_heading('三、个人微信方案（备选）', level=1)
h.runs[0].font.color.rgb = RGBColor(26, 95, 122)

doc.add_heading('3.1 什么是个人微信方案？', level=2)
doc.add_paragraph('在个人微信电脑上安装辅助软件，实现与企微方案类似的功能。')

p = doc.add_paragraph()
run = p.add_run('风险提示：')
run.bold = True
run.font.color.rgb = RGBColor(192, 57, 43)
p.add_run('个人微信方案使用第三方软件，存在封号风险，建议仅作为备选方案。')

doc.add_heading('3.2 支持功能', level=2)

table2 = doc.add_table(rows=6, cols=3)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
hdr2[0].text = '功能模块'
hdr2[1].text = '说明'
hdr2[2].text = '状态'
for cell in hdr2:
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

data2 = [
    ('自动欢迎语', '新客户添加后自动回复', '✓ 支持'),
    ('关键词回复', '收到特定关键词自动回复', '✓ 支持'),
    ('定时群发', '定时向客户推送内容', '✓ 支持'),
    ('群内自动回复', '群里@小助手自动回复', '✓ 支持'),
    ('消息管理', '客户消息记录存档', '✓ 支持'),
]
for i, (f, d, s) in enumerate(data2, 1):
    table2.rows[i].cells[0].text = f
    table2.rows[i].cells[1].text = d
    table2.rows[i].cells[2].text = s

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('说明：')
run.bold = True
p.add_run('个人微信平台限制，不支持自动批量添加好友功能。如需此功能，请选择企业微信方案。')

doc.add_heading('3.3 风险提示', level=2)
doc.add_paragraph('封号风险：使用第三方软件存在被封号的可能性', style='List Bullet')
doc.add_paragraph('维护成本：微信版本更新可能导致功能失效', style='List Bullet')
doc.add_paragraph('稳定性：不如官方API稳定', style='List Bullet')

doc.add_heading('3.4 适用场景', level=2)
doc.add_paragraph('已有个人微信客户资源', style='List Bullet')
doc.add_paragraph('愿意承担一定封号风险', style='List Bullet')
doc.add_paragraph('企业微信方案无法满足的特殊需求', style='List Bullet')

doc.add_page_break()

# 四、方案对比
h = doc.add_heading('四、方案对比', level=1)
h.runs[0].font.color.rgb = RGBColor(26, 95, 122)

table3 = doc.add_table(rows=6, cols=3)
table3.style = 'Table Grid'
hdr3 = table3.rows[0].cells
hdr3[0].text = '对比项'
hdr3[1].text = '企业微信方案'
hdr3[2].text = '个人微信方案'
for cell in hdr3:
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

data3 = [
    ('合规性', '✓ 官方合规，零风险', '⚠ 存在封号风险'),
    ('稳定性', '✓ 企业级稳定', '⚠ 可能受微信更新影响'),
    ('使用门槛', '需要企业微信（已认证更完整）', '个人微信即可'),
    ('核心功能', '✓ 自动加好友、欢迎语、群发等全部支持', '⚠ 部分功能受限（不支持自动加好友）'),
    ('维护成本', '✓ 低', '⚠ 需跟进微信版本'),
]
for i, (item, ent, pers) in enumerate(data3, 1):
    table3.rows[i].cells[0].text = item
    table3.rows[i].cells[1].text = ent
    table3.rows[i].cells[2].text = pers

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('推荐程度：企业微信方案 ⭐⭐⭐⭐⭐ 强烈推荐  |  个人微信方案 ⭐⭐ 备选')
run.bold = True

doc.add_page_break()

# 五、方案推荐
h = doc.add_heading('五、方案推荐', level=1)
h.runs[0].font.color.rgb = RGBColor(26, 95, 122)

doc.add_heading('5.1 强烈建议选择：企业微信方案', level=2)
p = doc.add_paragraph()
run = p.add_run('推荐理由：')
run.bold = True
run.font.color.rgb = RGBColor(26, 95, 122)

doc.add_paragraph('安全第一：企业微信是官方接口，不会封号，客户资源有保障', style='List Bullet')
doc.add_paragraph('稳定省心：一次部署，长期运行，无需担心微信更新', style='List Bullet')
doc.add_paragraph('功能够用：覆盖绝大多数自动化运营场景', style='List Bullet')

doc.add_heading('5.2 何时考虑个人微信方案？', level=2)
doc.add_paragraph('企业微信方案无法满足的特殊功能', style='List Bullet')
doc.add_paragraph('短期活动（可接受风险）', style='List Bullet')
doc.add_paragraph('技术验证/POC阶段', style='List Bullet')

doc.add_page_break()

# 六、常见问题
h = doc.add_heading('六、常见问题', level=1)
h.runs[0].font.color.rgb = RGBColor(26, 95, 122)

faqs = [
    ('Q：每天能发多少消息？',
     '企业微信限制：\n  · 群发消息：同一客户每天最多收到4条群发消息\n  · 系统内置频率控制，避免触发风控\n\n超出限制怎么办？\n  · 系统自动排队，分时段发送\n  · 内容多样化，避免触发风控\n  · 系统自动调控发送节奏'),
    ('Q：会被封号吗？',
     '· 企业微信方案：不会。使用官方API接口，完全合规。\n· 个人微信方案：存在一定封号风险。'),
    ('Q：需要准备什么？',
     '· 企业微信方案：企业微信账号\n· 个人微信方案：电脑、微信PC客户端、个人微信账号'),
    ('Q：企业微信方案是否需要认证？',
     '· 基础功能（欢迎语、关键词回复、定时群发）：已开通企业微信即可使用\n· 自动加好友（创建"联系我"二维码）：需要完成企业认证\n· 建议在采购前确认所需功能是否需要认证'),
    ('Q：数据存在哪里？',
     '本地服务器/云服务器，数据库由您自己掌控，不经过第三方。'),
    ('Q：能对接我们的现有系统吗？',
     '可以。系统支持两种对接方式：\n· 标准API接口：与现有CRM、ERP等系统对接\n· 智能助手：用自然语言指令控制自动化流程'),
    ('Q：智能助手是什么？',
     '智能助手是基于AI大模型的能力，可以用自然语言来操作微信自动化。无需学习复杂操作，说句话就能完成。'),
]

for q, a in faqs:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    run.font.color.rgb = RGBColor(26, 95, 122)
    doc.add_paragraph(a)

doc.add_page_break()

# 七、服务内容
h = doc.add_heading('七、服务内容', level=1)
h.runs[0].font.color.rgb = RGBColor(26, 95, 122)

table4 = doc.add_table(rows=5, cols=2)
table4.style = 'Table Grid'
hdr4 = table4.rows[0].cells
hdr4[0].text = '项目'
hdr4[1].text = '说明'
for cell in hdr4:
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

data4 = [
    ('系统部署', '在您的服务器上安装配置整套系统'),
    ('功能定制', '根据业务需求调整自动化流程'),
    ('培训支持', '提供使用培训和技术文档'),
    ('持续维护', '定期更新，保证系统稳定运行'),
]
for i, (item, desc) in enumerate(data4, 1):
    table4.rows[i].cells[0].text = item
    table4.rows[i].cells[1].text = desc

doc.add_paragraph()
doc.add_paragraph()

# 联系我们
h = doc.add_heading('联系我们', level=1)
h.runs[0].font.color.rgb = RGBColor(26, 95, 122)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('如有任何问题或需要进一步演示，欢迎随时联系')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(102, 102, 102)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('方案设计者：墨飞科技 | 本方案专为满足中小型企业微信运营需求设计')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(153, 153, 153)

doc.save('D:/my_project/ai_projects/wechat-automation-research/微信自动化运营方案-客户版.docx')
print('Word文档已生成成功！')
